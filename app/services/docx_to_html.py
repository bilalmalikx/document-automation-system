from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
import re
from pathlib import Path

class DocxToHtmlConverter:
    @staticmethod
    def convert(docx_path: Path) -> str:
        """Convert DOCX to HTML/CSS with full styling"""
        doc = Document(docx_path)
        
        html_parts = [
            '<!DOCTYPE html>',
            '<html>',
            '<head>',
            '<meta charset="UTF-8">',
            '<style>',
            '''
            body {
                font-family: 'Segoe UI', Arial, sans-serif;
                margin: 40px auto;
                padding: 20px;
                max-width: 900px;
                line-height: 1.6;
                color: #333;
                background: white;
            }
            h1 { color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }
            h2 { color: #34495e; margin-top: 25px; }
            h3 { color: #555; }
            table {
                border-collapse: collapse;
                width: 100%;
                margin: 20px 0;
                box-shadow: 0 1px 3px rgba(0,0,0,0.1);
            }
            th, td {
                border: 1px solid #ddd;
                padding: 12px;
                text-align: left;
            }
            th {
                background-color: #3498db;
                color: white;
                font-weight: bold;
            }
            tr:nth-child(even) { background-color: #f9f9f9; }
            ul, ol { margin: 15px 0; padding-left: 25px; }
            li { margin: 5px 0; }
            .bold { font-weight: bold; }
            .italic { font-style: italic; }
            .underline { text-decoration: underline; }
            .placeholder {
                background: #fff3cd;
                color: #856404;
                padding: 2px 6px;
                border-radius: 4px;
                font-family: monospace;
                display: inline-block;
            }
            table {
                width: 100%;
                border-collapse: collapse;
                margin: 20px 0;
            }
            th, td {
                border: 1px solid #ddd;
                padding: 12px;
                text-align: left;
                vertical-align: top;
            }
            th {
                background: linear-gradient(135deg, #3498db, #2980b9);
                color: white;
                font-weight: 600;
            }
            .company-name { font-size: 24px; font-weight: bold; color: #2c3e50; }
            .signature { margin-top: 50px; border-top: 2px solid #ccc; padding-top: 20px; }
            .footer { margin-top: 40px; text-align: center; font-size: 12px; color: #7f8c8d; }
            ''',
            '</style>',
            '</head>',
            '<body>'
        ]
        
        # Process paragraphs
        for para in doc.paragraphs:
            html_parts.append(DocxToHtmlConverter._process_paragraph(para))
        
        # Process tables
        for table in doc.tables:
            html_parts.append(DocxToHtmlConverter._process_table(table))
        
        html_parts.append('</body></html>')
        
        return '\n'.join(html_parts)
    
    @staticmethod
    def _process_paragraph(para: Paragraph) -> str:
        text = para.text
        if not text.strip():
            return ''
        
        # Check for placeholders and highlight them
        text = re.sub(
            r'\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}',
            r'<span class="placeholder">{{\1}}</span>',
            text
        )
        
        # Get paragraph style
        style = para.style.name.lower() if para.style else ''
        
        if 'heading 1' in style or 'title' in style:
            return f'<h1>{text}</h1>'
        elif 'heading 2' in style:
            return f'<h2>{text}</h2>'
        elif 'heading 3' in style:
            return f'<h3>{text}</h3>'
        elif 'list' in style or para.text.startswith('•') or para.text.startswith('-'):
            return f'<li>{text[1:].strip()}</li>'
        else:
            return f'<p>{text}</p>'
    
    @staticmethod
    def _process_table(table: Table) -> str:
        html = ['<table>']
        
        # Process header row
        if table.rows:
            html.append('<thead><tr>')
            for cell in table.rows[0].cells:
                html.append(f'<th>{cell.text}</th>')
            html.append('</tr></thead>')
        
        # Process body rows
        html.append('<tbody>')
        for row in table.rows[1:]:
            html.append('<tr>')
            for cell in row.cells:
                cell_text = re.sub(
                    r'\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}',
                    r'<span class="placeholder">{{\1}}</span>',
                    cell.text
                )
                html.append(f'<td>{cell_text}</td>')
            html.append('</tr>')
        html.append('</tbody>')
        
        html.append('</table>')
        return '\n'.join(html)