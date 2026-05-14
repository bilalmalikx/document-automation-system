from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session
from uuid import UUID, uuid4
from pathlib import Path
import shutil
import re
from datetime import datetime

from app.database.session import get_db
from app.models.template import Template
from app.models.template_field import TemplateField
from app.config import settings

router = APIRouter()


# ============ GET ALL TEMPLATES ============
@router.get("/")
def list_templates(db: Session = Depends(get_db)):
    templates = db.query(Template).filter(Template.is_active == True).all()
    return [
        {
            "id": str(t.id),
            "name": t.name,
            "description": t.description,
            "template_type": t.template_type
        }
        for t in templates
    ]


# ============ GET TEMPLATE SCHEMA (for form) ============
@router.get("/{template_id}/schema")
def get_template_schema(template_id: UUID, db: Session = Depends(get_db)):
    fields = db.query(TemplateField).filter(
        TemplateField.template_id == template_id
    ).order_by(TemplateField.display_order).all()
    
    template = db.query(Template).filter(Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    return {
        "template_id": template_id,
        "template_name": template.name,
        "fields": [
            {
                "name": f.placeholder_name,
                "label": f.field_label,
                "type": f.field_type,
                "required": f.is_required,
                "options": f.options
            }
            for f in fields
        ]
    }


# ============ GET TEMPLATE CONTENT (for Panel 1 & 2) ============

@router.get("/{template_id}/content")
def get_template_content(template_id: UUID, db: Session = Depends(get_db)):
    template = db.query(Template).filter(Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    content = ""
    
    if template.template_type == "html" and template.html_template_path:
        file_path = Path(template.html_template_path)
        if file_path.exists():
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
    
    elif template.template_type == "docx" and template.docx_template_path:
        file_path = Path(template.docx_template_path)
        if file_path.exists():
            try:
                # Try to import and use converter
                import sys
                sys.path.append(str(Path(__file__).parent.parent.parent))
                from app.services.docx_to_html import DocxToHtmlConverter
                content = DocxToHtmlConverter.convert(file_path)
                print(f"✅ Successfully converted DOCX to HTML")
            except ImportError as e:
                print(f"Import error: {e}")
                content = f"<p>Converter not found. Please create docx_to_html.py</p>"
            except Exception as e:
                print(f"Conversion error: {e}")
                content = f"<p>Error converting DOCX: {str(e)}</p>"
    
    return {
        "template_id": template_id,
        "name": template.name,
        "template_type": template.template_type,
        "content": content
    }


# ============ UPLOAD DOCX TEMPLATE ============
@router.post("/upload")
async def upload_template(
    name: str = Form(...),
    description: str = Form(None),
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only .docx files are allowed")
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_filename = f"{timestamp}_{file.filename}"
    template_path = Path(settings.TEMPLATE_DOCX_DIR) / unique_filename
    template_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Save file
    with open(template_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    # Extract placeholders
    placeholders = []
    try:
        from docx import Document
        doc = Document(template_path)
        
        # Extract text from all paragraphs
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            full_text.append(para.text)
        
        full_text_str = ' '.join(full_text)
        
        # Find all placeholders {{anything}}
        matches = re.findall(r'\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}', full_text_str)
        placeholders = list(dict.fromkeys(matches))
        
        print(f"✅ Found {len(placeholders)} placeholders: {placeholders}")
        
    except Exception as e:
        print(f"Error extracting placeholders: {e}")
        placeholders = []
    
    # Create template record
    template = Template(
        id=uuid4(),
        name=name,
        description=description,
        docx_template_path=f"app/templates/docx/{unique_filename}",
        template_type="docx",
        is_active=True
    )
    db.add(template)
    db.flush()
    
    # Delete old fields if any
    db.query(TemplateField).filter(TemplateField.template_id == template.id).delete()
    
    # Create template fields
    for idx, placeholder in enumerate(placeholders):
        label = placeholder.replace('_', ' ').title()
        
        # Guess field type
        field_type = "text"
        if "date" in placeholder.lower():
            field_type = "date"
        elif "email" in placeholder.lower():
            field_type = "email"
        elif "phone" in placeholder.lower():
            field_type = "tel"
        
        template_field = TemplateField(
            id=uuid4(),
            template_id=template.id,
            placeholder_name=placeholder,
            field_label=label,
            field_type=field_type,
            is_required=True,
            display_order=idx + 1
        )
        db.add(template_field)
    
    db.commit()
    
    return {
        "template_id": template.id,
        "name": name,
        "template_type": "docx",
        "fields_found": placeholders,
        "message": f"DOCX template uploaded successfully. Found {len(placeholders)} placeholders."
    }


# ============ UPLOAD HTML TEMPLATE ============
@router.post("/upload-html")
async def upload_html_template(
    name: str = Form(...),
    description: str = Form(None),
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    if not file.filename.endswith('.html'):
        raise HTTPException(status_code=400, detail="Only .html files are allowed")
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_filename = f"{timestamp}_{file.filename}"
    template_path = Path(settings.TEMPLATE_HTML_DIR) / unique_filename
    template_path.parent.mkdir(parents=True, exist_ok=True)
    
    with open(template_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    with open(template_path, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    placeholders = re.findall(r'\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}', html_content)
    unique_placeholders = list(dict.fromkeys(placeholders))
    
    template = Template(
        id=uuid4(),
        name=name,
        description=description,
        html_template_path=f"app/templates/html/{unique_filename}",
        template_type="html",
        is_active=True
    )
    db.add(template)
    db.flush()
    
    for idx, placeholder in enumerate(unique_placeholders):
        label = placeholder.replace('_', ' ').title()
        field_type = "text"
        if "date" in placeholder.lower():
            field_type = "date"
        elif "email" in placeholder.lower():
            field_type = "email"
        
        template_field = TemplateField(
            id=uuid4(),
            template_id=template.id,
            placeholder_name=placeholder,
            field_label=label,
            field_type=field_type,
            is_required=True,
            display_order=idx + 1
        )
        db.add(template_field)
    
    db.commit()
    
    return {
        "template_id": template.id,
        "name": name,
        "template_type": "html",
        "fields_found": unique_placeholders,
        "message": f"HTML template uploaded successfully. Found {len(unique_placeholders)} placeholders."
    }